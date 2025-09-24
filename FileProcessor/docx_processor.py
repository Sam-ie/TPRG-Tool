# docx_processor.py
import docx

class DocxProcessor:
    def read_file(self, path):
        # 读取docx文件
        doc = docx.Document(path)
        return [para.text for para in doc.paragraphs if para.text.strip()]

    def export_file(self, text, output_path):
        # 导出为docx文件
        doc = docx.Document()
        for line in text:
            if line.strip():  # 跳过空行
                doc.add_paragraph(line)
        doc.save(output_path)
        return True