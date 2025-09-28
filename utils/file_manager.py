import os
from typing import Optional, Tuple
from docx import Document
import pythoncom
import win32com.client


class FileManager:
    SUPPORTED_IMPORT_EXTENSIONS = ['.docx', '.doc', '.txt']
    SUPPORTED_EXPORT_EXTENSIONS = ['.docx', '.txt']

    def __init__(self, language_manager):
        self.language_manager = language_manager

    def read_file(self, file_path: str, progress_callback=None) -> Tuple[Optional[str], Optional[str]]:
        """读取文件内容，返回(内容, 错误信息)"""
        try:
            if not os.path.exists(file_path):
                return None, self.language_manager.get_text("file_not_exist")

            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.txt':
                return self._read_txt_file(file_path, progress_callback)
            elif ext == '.docx':
                return self._read_docx_file(file_path, progress_callback)
            elif ext == '.doc':
                return self._read_doc_file(file_path, progress_callback)
            else:
                return None, self.language_manager.get_text("unsupported_format")

        except Exception as e:
            return None, f"{self.language_manager.get_text('read_file_error')}: {str(e)}"

    def _read_txt_file(self, file_path: str, progress_callback=None) -> Tuple[Optional[str], Optional[str]]:
        """读取txt文件"""
        # 先获取文件行数来估算进度
        total_lines = 0
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                total_lines = sum(1 for _ in f)
        except:
            pass  # 如果无法用utf-8读取，后面会尝试其他编码

        # 尝试多种编码读取txt文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']

        for encoding in encodings:
            try:
                content_lines = []
                with open(file_path, 'r', encoding=encoding) as f:
                    for i, line in enumerate(f):
                        content_lines.append(line)
                        if progress_callback and total_lines > 0:
                            status = f"{self.language_manager.get_text('reading_file_progress')} ({encoding})"
                            progress_callback(i + 1, total_lines, status)

                content = ''.join(content_lines)
                return content, None
            except UnicodeDecodeError:
                continue
        return None, self.language_manager.get_text("decode_error")

    def _read_docx_file(self, file_path: str, progress_callback=None) -> Tuple[Optional[str], Optional[str]]:
        """读取docx文件"""
        try:
            doc = Document(file_path)
            paragraphs = doc.paragraphs
            total_paragraphs = len(paragraphs)

            content_lines = []
            for i, paragraph in enumerate(paragraphs):
                if paragraph.text.strip():  # 只处理非空段落
                    content_lines.append(paragraph.text)

                if progress_callback and total_paragraphs > 0:
                    progress_callback(i + 1, total_paragraphs, self.language_manager.get_text("parsing_docx"))

            content = '\n'.join(content_lines)
            return content, None
        except Exception as e:
            return None, f"{self.language_manager.get_text('read_docx_error')}: {str(e)}"

    def _read_doc_file(self, file_path: str, progress_callback=None) -> Tuple[Optional[str], Optional[str]]:
        """读取doc文件"""
        try:
            # 使用Word应用程序读取.doc文件
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)

            if progress_callback:
                progress_callback(50, 100, self.language_manager.get_text("converting_doc"))

            content = doc.Content.Text

            if progress_callback:
                progress_callback(100, 100, self.language_manager.get_text("conversion_complete"))

            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            return content, None
        except Exception as e:
            return None, f"{self.language_manager.get_text('read_doc_error')}: {str(e)}"

    def write_file(self, file_path: str, content: str, file_type: str) -> Optional[str]:
        """写入文件，返回错误信息（成功返回None）"""
        try:
            ext = os.path.splitext(file_path)[1].lower()

            if file_type == 'txt' or ext == '.txt':
                # 保存为txt文件
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return None

            elif file_type == 'docx' or ext == '.docx':
                # 保存为docx文件
                doc = Document()
                # 按行分割内容并添加到文档
                lines = content.split('\n')
                for line in lines:
                    if line.strip():  # 非空行才添加
                        doc.add_paragraph(line)
                doc.save(file_path)
                return None

            else:
                return self.language_manager.get_text("unsupported_export_format")

        except Exception as e:
            return f"{self.language_manager.get_text('save_file_error')}: {str(e)}"

    def is_supported_import_file(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in FileManager.SUPPORTED_IMPORT_EXTENSIONS

    def is_supported_export_file(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in FileManager.SUPPORTED_EXPORT_EXTENSIONS

    def get_file_filters(self, import_filter: bool = True) -> list:
        """获取文件过滤器"""
        if import_filter:
            return [
                (self.language_manager.get_text("supported_files"), "*.docx *.doc *.txt"),
                (self.language_manager.get_text("word_documents"), "*.docx"),
                (self.language_manager.get_text("word_97_2003_documents"), "*.doc"),
                (self.language_manager.get_text("text_files"), "*.txt"),
                (self.language_manager.get_text("all_files"), "*.*")
            ]
        else:
            return [
                (self.language_manager.get_text("word_documents"), "*.docx"),
                (self.language_manager.get_text("text_files"), "*.txt"),
                (self.language_manager.get_text("all_files"), "*.*")
            ]